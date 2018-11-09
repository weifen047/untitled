#encoding:utf-8

from docx import Document
import StringIO
#from docx.shared import Inches


document = Document('temp201810.docx')  # type: object

tables = document.tables




for cells in document.tables[3].columns[1].cells:
    print cells.text
    # for row in cells:
    #     print row[i]






document.save('temp201810-2.docx')





